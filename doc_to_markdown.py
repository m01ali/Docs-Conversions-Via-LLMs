import os
import argparse
import tempfile
import shutil
import base64
from pathlib import Path
import subprocess
import json
import re

# Import OpenAI for OpenRouter integration
from openai import OpenAI

# Import document processing libraries
import fitz  # PyMuPDF for PDF
from pdf2image import convert_from_path
import pytesseract 
from docx import Document
import pandas as pd
import openpyxl #excel
from bs4 import BeautifulSoup # HTML parsing
import markdown
import requests

class DocumentConverter:
    def __init__(self, openrouter_api_key, site_url="Your Website", site_name="Document Converter"):
        self.output_dir = "markdown_output"
        self.images_dir = os.path.join(self.output_dir, "images")
        self.temp_dir = tempfile.mkdtemp()
        
        # Set up OpenRouter client
        self.client = OpenAI(
            base_url="https://openrouter.ai/api/v1",
            api_key=openrouter_api_key,
        )
        self.site_url = site_url
        self.site_name = site_name
        
        # Initialize directories
        os.makedirs(self.output_dir, exist_ok=True)
        os.makedirs(self.images_dir, exist_ok=True)
    
    def __del__(self):
        # Clean up temp directory
        if hasattr(self, 'temp_dir') and os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)
    
    def convert_file(self, input_path):
        """Convert a file to Markdown based on its extension"""
        file_extension = os.path.splitext(input_path)[1].lower()
        output_filename = os.path.basename(input_path).rsplit('.', 1)[0] + '.md'
        output_path = os.path.join(self.output_dir, output_filename)
        
        print(f"Converting {input_path} to Markdown...")
        
        if self._try_markitdown_cli(input_path, output_path):
            print(f"Successfully converted using markitdown CLI to {output_path}")
            
            # Special handling for PDF images
            if input_path.lower().endswith('.pdf'):
                print("Adding PDF images to the converted markdown...")
                self._extract_images_from_pdf_post_markitdown(input_path, output_path)
            
            # Continue with normal processing
            self._extract_base64_images(output_path)
            self._fix_placeholder_image_references(output_path)
            self._enhance_with_llm(output_path)
            return output_path
        
        # If markitdown CLI fails or isn't suitable, use our custom conversion methods
        if file_extension in ['.pdf']:
            self._convert_pdf(input_path, output_path)
        elif file_extension in ['.docx', '.doc']:
            self._convert_docx(input_path, output_path)
        elif file_extension in ['.xlsx', '.xls']:
            self._convert_excel(input_path, output_path)
        elif file_extension in ['.html', '.htm']:
            self._convert_html(input_path, output_path)
        elif file_extension in ['.md', '.markdown', '.txt']:
            self._process_text(input_path, output_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
            
        # Post-process to enhance formatting
        self._enhance_with_llm(output_path)
        
        # After conversion is complete
        self._verify_image_paths(output_path)
            
        print(f"Successfully converted to {output_path}")
        return output_path
    
    def _extract_images_from_pdf_post_markitdown(self, pdf_path, markdown_path):
        """Extract images from PDF and add them to the markdown after MarkItDown conversion"""
        pdf_document = fitz.open(pdf_path)
        extracted_images = []
        
        # Read existing markdown content
        with open(markdown_path, 'r', encoding='utf-8') as f:
            markdown_content = f.read()
        
        # Extract images from all pages
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            image_list = page.get_images(full=True)
            
            if image_list:
                for img_index, img_info in enumerate(image_list):
                    xref = img_info[0]
                    base_img = pdf_document.extract_image(xref)
                    image_bytes = base_img["image"]
                    
                    # Save image
                    img_filename = f"image_p{page_num+1}_{img_index+1}.{base_img['ext']}"
                    img_path = os.path.join(self.images_dir, img_filename)
                    
                    try:
                        with open(img_path, "wb") as img_file:
                            img_file.write(image_bytes)
                        
                        print(f"Extracted image to {img_path}")
                        extracted_images.append((img_filename))
                    except Exception as e:
                        print(f"Error saving image: {e}")
            else:
                # If no vectorized images found, try rendering page as image
                pix = page.get_pixmap(alpha=False)
                img_filename = f"page_{page_num+1}.png"
                img_path = os.path.join(self.images_dir, img_filename)
                
                try:
                    pix.save(img_path)
                    print(f"Rendered page {page_num+1} as image: {img_path}")
                    extracted_images.append((img_filename))
                except Exception as e:
                    print(f"Error saving page image: {e}")
        
        # Append image references to the markdown content if any images were found
        if extracted_images:
            img_rel_path = f"images/{img_filename}".replace('\\', '/')
            markdown_content += f"![Image from document]({img_rel_path})\n\n"
            for img_filename in extracted_images:
                markdown_content += f"![Image from document](images/{img_filename})\n\n"
            
            # Write updated markdown
            with open(markdown_path, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
            
            print(f"Added {len(extracted_images)} image references to the markdown")
        
        pdf_document.close()
    
    
    def _try_markitdown_cli(self, input_path, output_path):
        """Attempt to use markitdown CLI tool if available"""
        if input_path.lower().endswith('.pdf'):
            return False

        
        try:
            # Check if markitdown CLI is installed and works for this file type
            result = subprocess.run(
                ["markitdown", input_path, "-o", output_path, "--keep-data-uris"],
                capture_output=True,
                text=True,
                check=False  # Don't raise exception on non-zero return code
            )
            return result.returncode == 0 and os.path.exists(output_path)
        except (subprocess.SubprocessError, FileNotFoundError):
            return False
    
    def _extract_images_from_pdf(self, pdf_path, page_num):
        """Extract images from PDF page"""
        image_list = []
        pdf_document = fitz.open(pdf_path)
        
        if page_num < len(pdf_document):
            page = pdf_document[page_num]
            image_list = page.get_images(full=True)
            
        image_paths = []
        for img_index, img_info in enumerate(image_list):
            xref = img_info[0]
            base_img = pdf_document.extract_image(xref)
            image_bytes = base_img["image"]
            
            # Save image
            img_filename = f"image_p{page_num+1}_{img_index+1}.{base_img['ext']}"
            img_path = os.path.join(self.images_dir, img_filename)
            with open(img_path, "wb") as img_file:
                img_file.write(image_bytes)
            
            image_paths.append((img_path, img_filename))
            
        pdf_document.close()
        return image_paths
    
    def _convert_pdf(self, pdf_path, output_path):
        """Enhanced PDF to Markdown conversion with robust image handling"""
        pdf_document = fitz.open(pdf_path)
        markdown_content = []
        image_references = []
        
        # Step 1: Extract all text for LLM processing
        print("Extracting text from PDF...")
        all_text = ""
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            text = page.get_text()
            if text.strip():
                all_text += text + "\n\n"
        
        # Step 2: Use LLM to format the text as markdown
        print("Converting text to markdown with LLM...")
        formatted_markdown = self._convert_with_llm(all_text, "pdf")
        if not formatted_markdown:
            formatted_markdown = all_text  # Fallback to raw text if LLM fails
        
        # Step 3: Extract and save images completely separately from text
        print("Extracting images from PDF...")
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            image_list = page.get_images(full=True)
            
            if image_list:
                for img_index, img_info in enumerate(image_list):
                    xref = img_info[0]
                    base_img = pdf_document.extract_image(xref)
                    image_bytes = base_img["image"]
                    
                    # Create unique image filename
                    img_filename = f"image_p{page_num+1}_{img_index+1}.{base_img['ext']}"
                    img_path = os.path.join(self.images_dir, img_filename)
                    
                    # Save the image
                    try:
                        with open(img_path, "wb") as img_file:
                            img_file.write(image_bytes)
                        print(f"Extracted image to {img_path}")
                        
                        # Create a markdown-compatible reference using RELATIVE path
                        rel_img_path = os.path.join("images", img_filename).replace("\\", "/")
                        image_references.append(f"![Image from page {page_num+1}]({rel_img_path})")
                    except Exception as e:
                        print(f"Error saving image: {e}")
            else:
                # If no vector images, try rendering page as image
                pix = page.get_pixmap(alpha=False)
                img_filename = f"page_{page_num+1}.png"
                img_path = os.path.join(self.images_dir, img_filename)
                
                try:
                    pix.save(img_path)
                    print(f"Rendered page {page_num+1} as image: {img_path}")
                    
                    # Create a markdown-compatible reference using RELATIVE path
                    rel_img_path = os.path.join("images", img_filename).replace("\\", "/")
                    image_references.append(f"![Page {page_num+1}]({rel_img_path})")
                except Exception as e:
                    print(f"Error saving page image: {e}")
        
        pdf_document.close()
        
        # Step 4: Create a test HTML file to verify image display
        html_test_path = os.path.join(os.path.dirname(output_path), "image_test.html")
        with open(html_test_path, 'w', encoding='utf-8') as f:
            f.write("<html><body>\n")
            for img_ref in image_references:
                img_path = img_ref.split('(')[1].split(')')[0]
                f.write(f'<p><img src="{img_path}" alt="Test image"></p>\n')
            f.write("</body></html>")
        
        # Step 5: Combine text and images in final markdown
        final_markdown = formatted_markdown
        if image_references:
            final_markdown += "\n\n## Document Images\n\n"
            final_markdown += "\n\n".join(image_references)
        
        # Write markdown content to file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(final_markdown)
        
        print(f"PDF successfully converted to {output_path} with {len(image_references)} images")
        print(f"To verify image paths, open {html_test_path} in a browser")
        
        return output_path

    def _convert_docx(self, docx_path, output_path):
        """Convert DOCX to Markdown with images and tables preserved"""
        doc = Document(docx_path)
        markdown_content = []
        
        # Process paragraphs and tables
        for i, element in enumerate(doc.element.body):
            if element.tag.endswith('p'):  # Paragraph
                paragraph = doc.paragraphs[i] if i < len(doc.paragraphs) else None
                if paragraph and paragraph.text.strip():
                    # Handle headings
                    if paragraph.style.name.startswith('Heading'):
                        level = paragraph.style.name[-1]
                        markdown_content.append(f"{'#' * int(level)} {paragraph.text}")
                    else:
                        markdown_content.append(paragraph.text)
            
            elif element.tag.endswith('tbl'):  # Table
                table_rows = []
                table = None
                for j, tbl in enumerate(doc.tables):
                    if j == i:
                        table = tbl
                        break
                
                if table:
                    # Extract headers
                    headers = [cell.text for cell in table.rows[0].cells]
                    table_rows.append('| ' + ' | '.join(headers) + ' |')
                    table_rows.append('| ' + ' | '.join(['---' for _ in headers]) + ' |')
                    
                    # Extract data rows
                    for row in table.rows[1:]:
                        row_data = [cell.text for cell in row.cells]
                        table_rows.append('| ' + ' | '.join(row_data) + ' |')
                    
                    markdown_content.append('\n'.join(table_rows))
        
        # Extract images
        img_index = 0
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                img_data = rel.target_part.blob
                img_ext = rel.target_ref.split('.')[-1]
                img_filename = f"image_{img_index}.{img_ext}"
                img_path = os.path.join(self.images_dir, img_filename)
                
                with open(img_path, "wb") as f:
                    f.write(img_data)
                
                relative_path = os.path.join("images", img_filename)
                markdown_content.append(f"\n\n![Image]({relative_path})\n\n")
                img_index += 1
        
        # Write markdown content to file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n\n'.join(markdown_content))
    
    def _convert_excel(self, excel_path, output_path):
        """Convert Excel to Markdown tables"""
        workbook = openpyxl.load_workbook(excel_path)
        markdown_content = []
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            markdown_content.append(f"## Sheet: {sheet_name}\n")
            
            # Determine the data range
            max_row = sheet.max_row
            max_col = sheet.max_column
            
            if max_row > 0 and max_col > 0:
                # Create table header
                header_row = []
                for col in range(1, max_col + 1):
                    cell_value = sheet.cell(row=1, column=col).value or ""
                    header_row.append(str(cell_value))
                
                markdown_content.append('| ' + ' | '.join(header_row) + ' |')
                markdown_content.append('| ' + ' | '.join(['---' for _ in range(max_col)]) + ' |')
                
                # Add data rows
                for row in range(2, max_row + 1):
                    row_data = []
                    for col in range(1, max_col + 1):
                        cell_value = sheet.cell(row=row, column=col).value or ""
                        row_data.append(str(cell_value))
                    markdown_content.append('| ' + ' | '.join(row_data) + ' |')
            
            markdown_content.append("\n\n")
        
        # Write markdown content to file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(markdown_content))
    
    def _convert_html(self, html_path, output_path):
        """Convert HTML to Markdown preserving structure"""
        with open(html_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Extract and save images
        img_index = 0
        for img in soup.find_all('img'):
            img_url = img.get('src', '')
            if img_url:
                if img_url.startswith('http'):
                    try:
                        response = requests.get(img_url, stream=True)
                        if response.status_code == 200:
                            img_ext = img_url.split('.')[-1].split('?')[0]
                            if len(img_ext) > 5:  # Not a valid extension
                                img_ext = 'png'
                            
                            img_filename = f"image_{img_index}.{img_ext}"
                            img_path = os.path.join(self.images_dir, img_filename)
                            
                            with open(img_path, 'wb') as f:
                                f.write(response.content)
                            
                            relative_path = os.path.join("images", img_filename)
                            img['src'] = relative_path
                            img_index += 1
                    except Exception as e:
                        print(f"Error downloading image {img_url}: {e}")
                elif img_url.startswith('data:image'):
                    # Handle base64 encoded images
                    try:
                        img_format = img_url.split(';')[0].split('/')[1]
                        img_data = img_url.split(',')[1]
                        img_filename = f"image_{img_index}.{img_format}"
                        img_path = os.path.join(self.images_dir, img_filename)
                        
                        with open(img_path, 'wb') as f:
                            f.write(base64.b64decode(img_data))
                        
                        relative_path = os.path.join("images", img_filename)
                        img['src'] = relative_path
                        img_index += 1
                    except Exception as e:
                        print(f"Error processing base64 image: {e}")
        
        # Convert HTML to markdown
        html_string = str(soup)
        
        # Use the LLM to convert HTML to proper markdown
        markdown_content = self._convert_with_llm(html_string, "html")
        
        # If LLM conversion fails, use a fallback method
        if not markdown_content:
            try:
                import html2text
                h = html2text.HTML2Text()
                h.ignore_links = False
                h.ignore_images = False
                h.ignore_tables = False
                markdown_content = h.handle(html_string)
            except ImportError:
                # Simple fallback if html2text is not available
                markdown_content = html_string
        
        # Write markdown content to file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(markdown_content)
    
    def _process_text(self, text_path, output_path):
        """Process text or markdown files, enhancing formatting if needed"""
        with open(text_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # For .md files, we may want to process and fix any issues
        if text_path.endswith(('.md', '.markdown')):
            # Copy images if they exist
            image_pattern = r'!\[.*?\]\((.*?)\)'
            for img_match in re.finditer(image_pattern, content):
                img_path = img_match.group(1)
                if os.path.exists(img_path):
                    img_filename = os.path.basename(img_path)
                    target_path = os.path.join(self.images_dir, img_filename)
                    shutil.copy(img_path, target_path)
                    
                    # Update the image reference in the content
                    relative_path = os.path.join("images", img_filename)
                    content = content.replace(img_path, relative_path)
        
        # Write content to file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)
    
    def _extract_base64_images(self, markdown_path):
        """Extract base64 images to files and update markdown references"""
        with open(markdown_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Find data URLs
        pattern = r'!\[([^\]]*)\]\(data:image/([^;]+);base64,([^)]+)\)'
        matches = list(re.finditer(pattern, content))
        
        if not matches:
            return  # No base64 images found
            
        print(f"Found {len(matches)} base64-encoded images to extract")
        
        def replace_match(match):
            alt_text = match.group(1)
            image_type = match.group(2)
            base64_data = match.group(3)
            
            # Generate image filename - FIX: Convert hash to string before slicing
            img_filename = f"image_{str(hash(base64_data))[:8]}.{image_type}"
            img_path = os.path.join(self.images_dir, img_filename)
            
            # Save image file
            try:
                with open(img_path, 'wb') as img_file:
                    img_file.write(base64.b64decode(base64_data))
                
                print(f"Extracted image to {img_path}")
                # Return updated markdown reference
                return f'![{alt_text}](images/{img_filename})'
            except Exception as e:
                print(f"Failed to extract image: {e}")
                return match.group(0)  # Return original if failed
        
        # Replace all data URLs
        updated_content = re.sub(pattern, replace_match, content)
        
        # Write updated markdown
        with open(markdown_path, 'w', encoding='utf-8') as f:
            f.write(updated_content)
        
    def _fix_placeholder_image_references(self, markdown_path):
        """Replace placeholder image references with actual image paths"""
        with open(markdown_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Check for placeholder patterns
        if 'path_to_image' in content:
            print("Found placeholder image references, replacing...")
            
            # Replace placeholders with sample images if available
            # You can add actual sample images to your images directory
            sample_img1 = os.path.join("images", "sample_image1.png")
            sample_img2 = os.path.join("images", "sample_image2.png")
            
            content = content.replace('(path_to_image1)', f'({sample_img1})')
            content = content.replace('(path_to_image2)', f'({sample_img2})')
            
            with open(markdown_path, 'w', encoding='utf-8') as f:
                f.write(content)    
        
    def _extract_tables_from_text(self, text):
        """Extract table-like structures from text"""
        # This is a simplified approach - in a real implementation, 
        # we would use more sophisticated pattern recognition
        tables = []
        lines = text.split('\n')
        current_table = []
        
        for line in lines:
            if '|' in line or '+--' in line or line.count('\t') >= 3:
                current_table.append(line)
            elif current_table:
                if len(current_table) >= 2:  # Minimum size for a table
                    tables.append('\n'.join(current_table))
                current_table = []
        
        if current_table and len(current_table) >= 2:
            tables.append('\n'.join(current_table))
        
        return tables
    
    def _verify_image_paths(self, markdown_path):
        """Verify image paths in markdown file and create HTML test file"""
        with open(markdown_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Extract image paths
        import re
        image_regex = r'!\[.*?\]\((.*?)\)'
        image_paths = re.findall(image_regex, content)
        
        if not image_paths:
            print("No image references found in markdown")
            return
        
        print(f"Found {len(image_paths)} image references:")
        for i, path in enumerate(image_paths):
            full_path = os.path.join(os.path.dirname(markdown_path), path)
            exists = os.path.exists(full_path)
            print(f"  {i+1}. {path} - {'EXISTS' if exists else 'MISSING'}")
        
        # Create a simple HTML file to test image display
        html_path = markdown_path.replace('.md', '_images.html')
        with open(html_path, 'w', encoding='utf-8') as f:
            f.write("<html><body>\n")
            f.write("<h1>Image Reference Test</h1>\n")
            for i, path in enumerate(image_paths):
                f.write(f'<h2>Image {i+1}</h2>\n')
                f.write(f'<p>Path: {path}</p>\n')
                f.write(f'<img src="{path}" style="max-width:800px;"><hr>\n')
            f.write("</body></html>")
        
        print(f"Created HTML test file: {html_path}")
        print(f"Open this file in a browser to check if images display correctly")
    
    
    def _convert_with_llm(self, content, source_format):
        """Use LLM to convert content from various formats to Markdown
        
        Args:
            content (str): The raw content to convert
            source_format (str): The format of the source content (e.g., "pdf", "html")
            
        Returns:
            str: Converted markdown content or None if conversion failed
        """
        try:
            # Get API key directly from command line args
            import sys
            api_key = None
            for i, arg in enumerate(sys.argv):
                if arg == '--api-key' and i+1 < len(sys.argv):
                    api_key = sys.argv[i+1]
                    break
            
            if not api_key:
                print("Error: No API key found. Please provide --api-key parameter.")
                return None
            
            # Customize prompt based on source format
            if source_format == "pdf":
                prompt = (
                    "You are a document conversion specialist. Your task is to convert PDF text to Markdown format WITHOUT summarizing or changing the text content in any way. Follow these rules strictly:\n"
                    "1. PRESERVE ALL ORIGINAL TEXT EXACTLY. Do not summarize, paraphrase, or omit any content.\n"
                    "2. Only add minimal Markdown formatting (## for obvious headings, * for bullet lists that already exist)\n"
                    "3. Preserve the exact paragraph structure as in the original\n"
                    "4. Keep all examples, code blocks, and technical details exactly as they appear\n"
                    "5. If you can't determine if something is a heading, keep it as plain text\n"
                    "\nHere is the PDF text to convert while preserving 100% of the original content:\n\n"
                )
            else:
                prompt = (
                    f"Convert the following {source_format} content to Markdown WITHOUT changing any text content.\n"
                    f"PRESERVE EVERY WORD EXACTLY as in the original. Do not summarize or paraphrase.\n"
                    f"Only add minimal Markdown formatting for structure.\n\n"
                    f"Content to convert (preserve all text exactly):\n\n"
                )
            
            
            # Handle content size limitations
            content_limit = 8000 if source_format == "pdf" else 10000
            truncated_content = content[:content_limit]
            prompt += truncated_content
            
            # Get site info from command line or use defaults
            site_url = "example.com"
            site_name = "Document Converter"
            
            for i, arg in enumerate(sys.argv):
                if arg == '--site-url' and i+1 < len(sys.argv):
                    site_url = sys.argv[i+1]
                elif arg == '--site-name' and i+1 < len(sys.argv):
                    site_name = sys.argv[i+1]
            
            # Create OpenAI client with OpenRouter API
            client = OpenAI(
                base_url="https://openrouter.ai/api/v1",
                api_key=api_key,
                default_headers={
                    "HTTP-Referer": site_url,
                    "X-Title": site_name
                }
            )
            
            # Send the conversion request to the API
            print(f"Sending content to LLM for {source_format} conversion...")
            completion = client.chat.completions.create(
                model="openai/gpt-4o",
                messages=[
                    {"role": "system", "content": "You are a document conversion specialist."},
                    {"role": "user", "content": prompt}
                ]
            )
            
            # Extract and return the converted markdown
            converted_text = completion.choices[0].message.content
            print(f"Successfully converted {len(truncated_content)} characters with LLM")
            return converted_text
            
        except Exception as e:
            print(f"Error using LLM for conversion: {e}")
            return None

    def _enhance_with_llm(self, markdown_path):
        """Enhance the converted markdown with LLM"""
        try:
            with open(markdown_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Only process if the content isn't too long
            if len(content) > 100000:
                print("Content too long for LLM enhancement, skipping...")
                return
                
            prompt = """
            Review and enhance the following markdown content:
            1. Fix any formatting issues
            2. Ensure tables are properly formatted
            3. Make sure image references are correct
            4. Preserve the original document structure
            5. Do NOT add any new content or commentary

            Markdown content:
            """
            
            # Process in chunks if needed
            chunk_size = 8000
            chunks = [content[i:i+chunk_size] for i in range(0, len(content), chunk_size)]
            enhanced_content = []
            
            for i, chunk in enumerate(chunks):
                completion = self.client.chat.completions.create(
                    extra_headers={
                        "HTTP-Referer": self.site_url, 
                        "X-Title": self.site_name,
                    },
                    model="openai/gpt-4o",
                    messages=[
                        {
                            "role": "system",
                            "content": "You are a document formatting specialist. Fix markdown formatting issues without changing content."
                        },
                        {
                            "role": "user",
                            "content": f"{prompt}\n\n{chunk}"
                        }
                    ]
                )
                
                enhanced_chunk = completion.choices[0].message.content
                enhanced_content.append(enhanced_chunk)
            
            # Combine enhanced chunks and write back to file
            enhanced_full = '\n'.join(enhanced_content)
            
            # Strip any added markdown code fences if present
            enhanced_full = re.sub(r'^```markdown\s*', '', enhanced_full)
            enhanced_full = re.sub(r'\s*```$', '', enhanced_full)
            
            with open(markdown_path, 'w', encoding='utf-8') as f:
                f.write(enhanced_full)
                
        except Exception as e:
            print(f"Error enhancing markdown: {e}")

def main():
    parser = argparse.ArgumentParser(description='Convert documents to Markdown with preserved formatting')
    parser.add_argument('input_file', help='Input document file (PDF, DOCX, XLSX, HTML, MD)')
    parser.add_argument('--api-key', required=True, help='OpenRouter API key')
    parser.add_argument('--site-url', default="https://example.com", help='Your site URL for OpenRouter')
    parser.add_argument('--site-name', default="Document Converter", help='Your site name for OpenRouter')
    
    args = parser.parse_args()
    
    converter = DocumentConverter(
        openrouter_api_key=args.api_key,
        site_url=args.site_url,
        site_name=args.site_name
    )
    
    output_path = converter.convert_file(args.input_file)
    print(f"Document successfully converted to {output_path}")

if __name__ == "__main__":
    main()